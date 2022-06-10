from docx.api import Document
import openpyxl
import datetime
import time
import tabula
import pandas as pd
import pandas
import xlsxwriter
from openpyxl import load_workbook
import os


def generate_time_stamp():
    ts = time.time()
    st = datetime.datetime.fromtimestamp(ts).strftime('%Y_%m_%d_%H_%M_%S')
    return st


# def delete_generated_excel(file_path):
#     if os.path.exists(file_path):
#         os.remove(file_path)
#     else:
#         print("The file does not exist")


def convert_docx_to_excel_4D29(docx_path):
    document = Document(docx_path)
    tables = document.tables
    df = pd.DataFrame()

    for table in document.tables:
        print("###############################################################")

        for row in table.rows:

            text = [cell.text for cell in row.cells]
            print("printing cell text")
            print(text)

            df = df.append([text], ignore_index=True)
            print("printing data frame")
            print(df)

    #storing the excel file with time stamp in a docx_excel_path

    docx_excel_path = "test_docx_4D29_" + generate_time_stamp() + ".xlsx"
    print("@@@@@@@@@@@@@@@@ printing excel file name @@@@@@@@@@@@@@@@@@@@@@@@@")
    print(docx_excel_path)

    df.to_excel(docx_excel_path)

    wb = openpyxl.load_workbook(docx_excel_path)
    sheet = wb['Sheet1']
    print(sheet['A1'].value)
    status = sheet['A1'].value
    print(status)
    return docx_excel_path


def convert_pdf_to_excel_4D29(pdf_path):
    df = tabula.read_pdf(pdf_path, pages='all')
    excel_path = "test_pdf_4D29_" + str(generate_time_stamp()) + ".xlsx"
    workbook = xlsxwriter.Workbook(excel_path)
    workbook.close()
    book = load_workbook(excel_path)
    writer = pandas.ExcelWriter(excel_path, engine='openpyxl')
    writer.book = book
    writer.sheets = {ws.title: ws for ws in book.worksheets}
    for i in range(0, len(df), 2):
        for sheetname in writer.sheets:
            df[i].to_excel(writer, sheet_name=sheetname, startrow=writer.sheets[sheetname].max_row, index=False,
                           header=False)

    writer.save()
    return excel_path


def compare_generated_excel_4D29(docx_file_path, pdf_file_path):
    docx_path = convert_docx_to_excel_4D29(docx_file_path)
    pdf_path = convert_pdf_to_excel_4D29(pdf_file_path)
    df_docx = pd.read_excel(docx_path)
    df_pdf = pd.read_excel(pdf_path)
    row_count_pdf = len(df_pdf)
    col_count_pdf = len(df_pdf.columns)
    data_generated_file_pdf = {"Seq.Number": [], "SampleName": [], "CPM": [], "FileName": [], "Analyte": []}
    dict_list_generated = []
    for key in data_generated_file_pdf.keys():
        dict_list_generated.append(key)
    for j in range(col_count_pdf):
        list_key = dict_list_generated[j]
        data_list = data_generated_file_pdf.get(list_key)
        for i in range(row_count_pdf):
            data_list.append(df_pdf.iloc[i, j])
    for k in range(len(dict_list_generated)):
        data_list = data_generated_file_pdf.get(dict_list_generated[k])
        column_name = data_list[0]
        data_list.pop(0)
        if column_name in data_list:
            data_list.remove(column_name)
            if column_name in data_list:
                data_list.remove(column_name)
    # print(data_generated_file_pdf)
    row_count_docx = len(df_pdf)
    col_count_docx = len(df_pdf.columns)
    data_generated_file_docx = {"Seq.Number": [], "SampleName": [], "CPM": [], "FileName": [], "Analyte": []}
    dict_list_generated_docx = []
    for key in data_generated_file_docx.keys():
        dict_list_generated_docx.append(key)
    for j in range(col_count_docx):
        list_key = dict_list_generated_docx[j]
        data_list = data_generated_file_docx.get(list_key)
        for i in range(row_count_docx - 2):
            data_list.append(df_docx.loc[i, j])
    for k in range(len(dict_list_generated_docx)):
        data_list = data_generated_file_docx.get(dict_list_generated_docx[k])
        data_list.pop(0)
    # print(data_generated_file_docx)
    del data_generated_file_pdf['FileName']
    del data_generated_file_docx['FileName']
    # print(data_generated_file_docx)
    # print(data_generated_file_pdf)
    key_list = []
    for key in data_generated_file_docx.keys():
        key_list.append(key)
    for l in range(len(key_list)):
        data_generated_value = data_generated_file_pdf.get(key_list[l])
        data_actual_value = data_generated_file_docx.get(key_list[l])
        for n in range(len(data_actual_value)):
            if str(data_actual_value[n]).strip() == str(data_generated_value[n]).strip():
                print(
                    "The values present in the DOCX file {actual_value} is matching with the value of Generated PDF {generated_value}".format
                    (actual_value=data_actual_value[n], generated_value=data_generated_value[n]))
            else:
                print(
                    (
                        "The values present in the DOCX file {actual_value} is not matching with the value of Generated PDF {generated_value}").format(
                        actual_value=data_actual_value[n], generated_value=data_generated_value[n]))
    delete_generated_excel(docx_path)
    delete_generated_excel(pdf_path)


# compare_generated_excel_4D29('DOCX_FILE/Test 4D29 Step 1.docx','PDF_FILE/Test 4D29 Step 1 generated.pdf')
compare_generated_excel_4D29('Test 4D29 Step 1.docx','Test 4D29 Step 1.pdf')