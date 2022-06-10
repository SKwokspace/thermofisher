from docx import Document

doc = Document(r'C:\Users\MohanMadhuri\PycharmProjects\thermofisher\newfile.docx')
# print('The whole content of the document:->>>')
for para in doc.paragraphs:
    print(para.text)
# for para in doc.paragraphs:
#     print(para.text)